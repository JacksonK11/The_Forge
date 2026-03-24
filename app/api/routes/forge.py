"""
app/api/routes/forge.py
Blueprint submission, build management, and update routes.

POST /forge/submit                                — submit blueprint text, queue build
POST /forge/submit-file                           — submit blueprint file (.docx/.pdf/.py/.txt etc)
POST /forge/submit-with-files                     — submit blueprint text + attached files, queue build
POST /forge/runs/{id}/approve                     — approve parsed spec, start code generation
POST /forge/runs/{id}/regenerate/{file_path}      — regenerate single file
GET  /forge/runs/{id}/package                     — download completed ZIP
POST /forge/update                                — queue a targeted codebase update
GET  /forge/updates                               — list all update runs
GET  /forge/updates/{update_id}                   — get update run detail
"""

import io
import re
import uuid
from io import BytesIO
from typing import List, Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import JSONResponse, StreamingResponse
from loguru import logger
from pydantic import BaseModel, Field
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.ratelimit import limiter
from app.api.services.error_translator import translate_error_for_storage
from memory.database import get_db
from memory.models import FileStatus, ForgeFile, ForgeRun, ForgeUpdate, RunStatus
from pipeline.pipeline import run_pipeline_sync
router = APIRouter()

# Binary document extensions that require special parsers
_BINARY_DOC_EXTENSIONS = {".docx", ".pdf"}

# All extensions that map to explicit UTF-8 plain-text treatment (no special parser)
# Any other extension also falls through to UTF-8 attempt — this list is informational.
_PLAIN_TEXT_EXTENSIONS = {
    ".py", ".txt", ".toml", ".json", ".md", ".yml", ".yaml",
    ".js", ".ts", ".jsx", ".tsx", ".html", ".css", ".scss", ".less",
    ".rs", ".go", ".sh", ".bash", ".zsh", ".fish",
    ".cfg", ".ini", ".env", ".sql",
    ".r", ".rb", ".php", ".swift", ".kt", ".dart",
    ".vue", ".svelte", ".prisma", ".graphql", ".proto",
    ".makefile", ".dockerfile", ".xml", ".csv", ".tsv",
    ".java", ".c", ".cpp", ".h", ".hpp", ".cs",
    ".tf", ".hcl", ".nix",
}


# ── Pydantic models ──────────────────────────────────────────────────────────


class SubmitBlueprintRequest(BaseModel):
    title: str = Field(..., min_length=1, max_length=255)
    blueprint_text: str = Field(..., min_length=50)
    repo_name: Optional[str] = None
    push_to_github: bool = True


class SubmitBlueprintResponse(BaseModel):
    run_id: str
    status: str
    message: str


class ApproveSpecResponse(BaseModel):
    run_id: str
    status: str
    message: str


class ResumeRunResponse(BaseModel):
    run_id: str
    status: str
    message: str


class RegenerateFileResponse(BaseModel):
    run_id: str
    file_path: str
    status: str
    message: str


class UpdateRequest(BaseModel):
    github_repo_url: str = Field(..., min_length=1, max_length=500)
    change_description: str = Field(..., min_length=10)
    title: Optional[str] = None
    callback_url: Optional[str] = None


class UpdateResponse(BaseModel):
    update_id: str
    status: str
    message: str


class FileExtractResponse(BaseModel):
    text: str
    filename: str


class SubmitWithFilesResponse(BaseModel):
    run_id: str
    status: str
    files_attached: int = 0


# ── Routes ────────────────────────────────────────────────────────────────────


@router.post("/submit", response_model=SubmitBlueprintResponse)
@limiter.limit("10/hour")
async def submit_blueprint(
    request: Request,
    body: SubmitBlueprintRequest,
    session: AsyncSession = Depends(get_db),
) -> SubmitBlueprintResponse:
    """
    Submit a blueprint for processing.
    Creates a ForgeRun record and queues the pipeline job.
    Returns the run_id for tracking.
    """
    from app.api.main import get_build_queue

    run_id = str(uuid.uuid4())
    resolved_repo_name = body.repo_name or _slug(body.title)

    run = ForgeRun(
        run_id=run_id,
        title=body.title,
        blueprint_text=body.blueprint_text,
        status=RunStatus.QUEUED.value,
        repo_name=resolved_repo_name,
        push_to_github=body.push_to_github,
    )
    session.add(run)
    await session.commit()

    try:
        queue = get_build_queue()
        queue.enqueue(
            run_pipeline_sync,
            run_id,
            job_id=f"build-{run_id}",
            job_timeout=3600,
        )
        logger.info(
            f"Build queued: run_id={run_id} title='{body.title}' "
            f"repo={resolved_repo_name} push_to_github={body.push_to_github}"
        )
    except Exception as exc:
        run.status = RunStatus.FAILED.value
        run.error_message = translate_error_for_storage(f"Failed to queue build: {exc}")
        await session.commit()
        logger.error(f"Failed to queue build {run_id}: {exc}")
        raise HTTPException(status_code=500, detail=f"Failed to queue build: {exc}")

    return SubmitBlueprintResponse(
        run_id=run_id,
        status=RunStatus.QUEUED.value,
        message="Blueprint accepted. Build queued.",
    )


@router.post("/submit-with-files", response_model=SubmitWithFilesResponse)
@limiter.limit("10/hour")
async def submit_with_files(
    request: Request,
    title: str = Form(...),
    blueprint_text: str = Form(...),
    repo_name: Optional[str] = Form(None),
    push_to_github: bool = Form(True),
    files: List[UploadFile] = File(default=[]),
    session: AsyncSession = Depends(get_db),
) -> SubmitWithFilesResponse:
    """
    Submit blueprint text with attached files for processing.

    Accepts multipart form data with:
    - title: Build title
    - blueprint_text: The user's typed instructions/prompt
    - repo_name: Optional repository name
    - push_to_github: Whether to push to GitHub (default True)
    - files: Zero or more attached files (any type)

    For each uploaded file, extracts text content server-side using file_extractor.
    Combines blueprint_text + all extracted file contents into one combined blueprint.
    Creates a ForgeRun and queues the pipeline job with the combined text.

    Returns {run_id, status: "queued"}.
    """
    from app.api.main import get_build_queue

    # Extract text from each attached file with smart capping.
    # Priority: code files (.py, .toml, .yml, .js, .ts) > plain text > .docx
    # Cap: 80 lines or 8000 chars per file, 25000 chars total across all files.
    _CODE_EXTS = {".py", ".toml", ".yml", ".yaml", ".js", ".ts", ".jsx", ".tsx",
                  ".go", ".rs", ".sh", ".sql", ".json", ".env", ".cfg", ".ini"}
    _DOC_EXTS = {".docx", ".pdf"}

    raw_extractions: list[tuple[str, str, int]] = []  # (filename, text, priority)
    skipped_files: list[str] = []

    for upload_file in files:
        filename = upload_file.filename or "unknown"
        ext = _get_file_extension(filename)
        try:
            file_bytes = await upload_file.read()
        except Exception as exc:
            logger.error(f"Failed to read uploaded file '{filename}': {exc}")
            skipped_files.append(filename)
            continue

        try:
            raw_text = _extract_text_from_file(file_bytes, filename)
        except Exception as exc:
            logger.warning(f"Could not extract '{filename}' — skipping: {exc}")
            skipped_files.append(filename)
            continue

        # For DOCX: extract only headings and first 2 sentences per section
        if ext == ".docx":
            raw_text = _extract_docx_outline(raw_text)

        # Cap per-file: 80 lines or 8000 chars (whichever is smaller)
        lines = raw_text.splitlines()
        if len(lines) > 80:
            raw_text = "\n".join(lines[:80]) + f"\n[... {len(lines)-80} more lines truncated]"
        if len(raw_text) > 8000:
            raw_text = raw_text[:8000] + f"\n[... truncated at 8000 chars]"

        priority = 0 if ext in _CODE_EXTS else (2 if ext in _DOC_EXTS else 1)
        raw_extractions.append((filename, raw_text, priority))
        logger.info(
            f"Extracted {len(raw_text)} chars from '{filename}' "
            f"(ext={ext}, priority={priority}, raw_bytes={len(file_bytes)})"
        )

    # Sort by priority (code files first, docs last) and cap total at 25000 chars
    raw_extractions.sort(key=lambda x: x[2])
    file_sections: list[str] = []
    total_file_chars = 0
    TOTAL_FILE_CAP = 25_000

    for filename, text, _ in raw_extractions:
        remaining = TOTAL_FILE_CAP - total_file_chars
        if remaining <= 0:
            logger.warning(f"Total file cap reached — skipping '{filename}'")
            skipped_files.append(filename)
            continue
        if len(text) > remaining:
            text = text[:remaining] + "\n[... truncated to fit total cap]"
        file_sections.append(f"=== ATTACHED FILE: {filename} ===\n{text}")
        total_file_chars += len(text)

    if skipped_files:
        logger.warning(f"Skipped {len(skipped_files)} files: {skipped_files}")

    logger.info(
        f"Attached files: {len(file_sections)} processed, "
        f"{total_file_chars:,} chars total (cap={TOTAL_FILE_CAP:,})"
    )

    # Combine blueprint text with extracted file contents
    if file_sections:
        combined_blueprint = blueprint_text + "\n\n" + "\n\n".join(file_sections)
    else:
        combined_blueprint = blueprint_text

    # Validate combined blueprint has enough content
    if len(combined_blueprint.strip()) < 50:
        raise HTTPException(
            status_code=400,
            detail="Combined blueprint text is too short (minimum 50 characters). "
            "Add more detail to your instructions or attach files with content.",
        )

    run_id = str(uuid.uuid4())
    resolved_repo_name = repo_name or _slug(title)

    run = ForgeRun(
        run_id=run_id,
        title=title,
        blueprint_text=combined_blueprint,
        status=RunStatus.QUEUED.value,
        repo_name=resolved_repo_name,
        push_to_github=push_to_github,
    )
    session.add(run)
    await session.commit()

    try:
        queue = get_build_queue()
        queue.enqueue(
            run_pipeline_sync,
            run_id,
            job_id=f"build-{run_id}",
            job_timeout=3600,
        )
        logger.info(
            f"Build with files queued: run_id={run_id} title='{title}' "
            f"repo={resolved_repo_name} files={len(files)} "
            f"combined_chars={len(combined_blueprint)} push_to_github={push_to_github}"
        )
    except Exception as exc:
        run.status = RunStatus.FAILED.value
        run.error_message = translate_error_for_storage(f"Failed to queue build: {exc}")
        await session.commit()
        logger.error(f"Failed to queue build with files {run_id}: {exc}")
        raise HTTPException(status_code=500, detail=f"Failed to queue build: {exc}")

    return SubmitWithFilesResponse(
        run_id=run_id,
        status=RunStatus.QUEUED.value,
        files_attached=len(files),
    )


@router.post("/submit-file")
async def submit_blueprint_file(
    file: UploadFile = File(...),
    title: Optional[str] = Form(None),
    repo_name: Optional[str] = Form(None),
    push_to_github: bool = Form(True),
    session: AsyncSession = Depends(get_db),
) -> JSONResponse:
    """
    Upload a file and extract its text content.

    For .docx files: uses python-docx to extract paragraph text.
    For .pdf files: uses pypdf to extract page text.
    For .py, .txt, .toml, .json, .md, .yml files: decodes as UTF-8.

    Returns JSON: {"text": extracted_content, "filename": original_filename}

    If a title is provided, also creates a ForgeRun and queues a build
    (legacy behaviour for direct file-to-build submissions).
    """
    filename = file.filename or "unknown"

    try:
        file_bytes = await file.read()
    except Exception as exc:
        logger.error(f"Failed to read uploaded file '{filename}': {exc}")
        raise HTTPException(status_code=400, detail=f"Failed to read file: {exc}")

    try:
        extracted_text = _extract_text_from_file(file_bytes, filename)
    except ValueError as exc:
        logger.error(f"Failed to extract text from '{filename}': {exc}")
        raise HTTPException(status_code=400, detail=str(exc))
    except Exception as exc:
        logger.error(f"Unexpected error extracting text from '{filename}': {exc}")
        raise HTTPException(status_code=400, detail=f"Could not parse file: {exc}")

    logger.info(
        f"File text extracted: filename='{filename}' "
        f"chars={len(extracted_text)} bytes={len(file_bytes)}"
    )

    # If no title provided, just return the extracted text (frontend text-extraction mode)
    if not title:
        return JSONResponse(
            content={"text": extracted_text, "filename": filename}
        )

    # Legacy mode: title provided → create a ForgeRun and queue build
    if len(extracted_text.strip()) < 50:
        raise HTTPException(
            status_code=400,
            detail="Extracted blueprint text is too short. Check the file content.",
        )

    from app.api.main import get_build_queue

    run_id = str(uuid.uuid4())
    resolved_repo_name = repo_name or _slug(title)

    run = ForgeRun(
        run_id=run_id,
        title=title,
        blueprint_text=extracted_text,
        blueprint_file_path=filename,
        status=RunStatus.QUEUED.value,
        repo_name=resolved_repo_name,
        push_to_github=push_to_github,
    )
    session.add(run)
    await session.commit()

    try:
        queue = get_build_queue()
        queue.enqueue(
            run_pipeline_sync,
            run_id,
            job_id=f"build-{run_id}",
            job_timeout=3600,
        )
    except Exception as exc:
        run.status = RunStatus.FAILED.value
        run.error_message = translate_error_for_storage(f"Failed to queue build: {exc}")
        await session.commit()
        logger.error(f"Failed to queue file build {run_id}: {exc}")
        raise HTTPException(status_code=500, detail=f"Failed to queue build: {exc}")

    logger.info(
        f"File build queued: run_id={run_id} file='{filename}' "
        f"repo={resolved_repo_name} push_to_github={push_to_github}"
    )

    return JSONResponse(
        content={
            "text": extracted_text,
            "filename": filename,
            "run_id": run_id,
            "status": RunStatus.QUEUED.value,
            "message": "Blueprint file accepted. Build queued.",
        }
    )


@router.post("/runs/{run_id}/approve", response_model=ApproveSpecResponse)
async def approve_spec(
    run_id: str,
    session: AsyncSession = Depends(get_db),
) -> ApproveSpecResponse:
    """
    Approve the parsed spec and continue with code generation.
    Called from the dashboard Spec Confirmation screen (Stage 3).
    Run must be in CONFIRMING status.
    """
    from app.api.main import get_build_queue

    result = await session.execute(
        select(ForgeRun).where(ForgeRun.run_id == run_id)
    )
    run = result.scalar_one_or_none()
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")
    if run.status != RunStatus.CONFIRMING.value:
        raise HTTPException(
            status_code=400,
            detail=f"Run is in status '{run.status}', not 'confirming'. Cannot approve.",
        )

    run.status = RunStatus.ARCHITECTING.value
    await session.commit()

    queue = get_build_queue()
    queue.enqueue(
        run_pipeline_sync,
        run_id,
        "resume_from_architecture",
        job_id=f"build-{run_id}-arch",
        job_timeout=3600,
    )

    logger.info(f"Spec approved, resuming pipeline: run_id={run_id}")
    return ApproveSpecResponse(
        run_id=run_id,
        status=RunStatus.ARCHITECTING.value,
        message="Spec approved. Generating architecture and starting code generation.",
    )


@router.post("/runs/{run_id}/resume", response_model=ResumeRunResponse)
async def resume_run(
    run_id: str,
    session: AsyncSession = Depends(get_db),
) -> ResumeRunResponse:
    """
    Re-queue a stalled or killed pipeline job to resume code generation.
    Only valid for runs in 'generating' or 'architecting' status.
    Skips already-complete files and continues from where the job died.
    """
    from app.api.main import get_build_queue

    result = await session.execute(select(ForgeRun).where(ForgeRun.run_id == run_id))
    run = result.scalar_one_or_none()
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")
    resumable = (RunStatus.GENERATING.value, RunStatus.ARCHITECTING.value, RunStatus.FAILED.value)
    if run.status not in resumable:
        raise HTTPException(
            status_code=400,
            detail=f"Run is in status '{run.status}'. Resume only valid for generating, architecting, or failed.",
        )

    # For failed runs, resume based on how far they got (manifest present = was generating)
    if run.status == RunStatus.FAILED.value:
        resume_from = "generating" if run.manifest_json else "resume_from_architecture"
        run.status = RunStatus.GENERATING.value if run.manifest_json else RunStatus.ARCHITECTING.value
        await session.commit()
    else:
        resume_from = "generating" if run.status == RunStatus.GENERATING.value else "resume_from_architecture"

    queue = get_build_queue()
    queue.enqueue(
        run_pipeline_sync,
        run_id,
        resume_from,
        job_id=f"build-{run_id}-resume",
        job_timeout=7200,
    )

    logger.info(f"[{run_id}] Pipeline re-queued for resume from '{resume_from}'")
    return ResumeRunResponse(
        run_id=run_id,
        status=run.status,
        message=f"Pipeline re-queued. Resuming from '{resume_from}', skipping completed files.",
    )


@router.post("/runs/{run_id}/regenerate/{file_path:path}", response_model=RegenerateFileResponse)
async def regenerate_file(
    run_id: str,
    file_path: str,
    session: AsyncSession = Depends(get_db),
) -> RegenerateFileResponse:
    """
    Regenerate a single file from a completed or failed run.
    Useful for iterating on a specific file without rebuilding the entire agent.
    Cost: ~£0.02–0.04 per file.
    """
    from app.api.main import get_build_queue
    from pipeline.pipeline import regenerate_file_sync

    result = await session.execute(
        select(ForgeRun).where(ForgeRun.run_id == run_id)
    )
    run = result.scalar_one_or_none()
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")
    if run.status not in (RunStatus.COMPLETE.value, RunStatus.FAILED.value):
        raise HTTPException(
            status_code=400,
            detail="Can only regenerate files from completed or failed runs",
        )

    file_result = await session.execute(
        select(ForgeFile).where(
            ForgeFile.run_id == run_id, ForgeFile.file_path == file_path
        )
    )
    forge_file = file_result.scalar_one_or_none()
    if not forge_file:
        raise HTTPException(status_code=404, detail="File not found in this run")

    queue = get_build_queue()
    queue.enqueue(
        regenerate_file_sync,
        run_id,
        file_path,
        job_id=f"regen-{run_id}-{file_path[:30]}",
        job_timeout=300,
    )

    logger.info(f"File regeneration queued: run_id={run_id} file={file_path}")
    return RegenerateFileResponse(
        run_id=run_id,
        file_path=file_path,
        status="queued",
        message="File regeneration queued.",
    )


@router.get("/runs/{run_id}/package")
async def download_package(
    run_id: str,
    session: AsyncSession = Depends(get_db),
) -> StreamingResponse:
    """
    Download the completed ZIP package for a build.
    Run must be in COMPLETE status and package_path must be set.
    """
    result = await session.execute(
        select(ForgeRun).where(ForgeRun.run_id == run_id)
    )
    run = result.scalar_one_or_none()
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")
    if run.status != RunStatus.COMPLETE.value:
        raise HTTPException(
            status_code=400, detail=f"Run is not complete (status: {run.status})"
        )
    if not run.package_data and not run.package_path:
        raise HTTPException(status_code=404, detail="Package not found for this run")

    if run.package_data:
        package_bytes = run.package_data
    else:
        try:
            with open(run.package_path, "rb") as f:
                package_bytes = f.read()
        except FileNotFoundError:
            raise HTTPException(status_code=404, detail="Package file missing from disk")

    agent_slug = run.spec_json.get("agent_slug", "agent") if run.spec_json else "agent"
    filename = f"{agent_slug}-forge-package.zip"

    return StreamingResponse(
        io.BytesIO(package_bytes),
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Update endpoints ──────────────────────────────────────────────────────────


@router.post("/update", response_model=UpdateResponse)
async def submit_update(
    request: UpdateRequest,
    session: AsyncSession = Depends(get_db),
) -> UpdateResponse:
    """
    Queue a targeted codebase update for an existing GitHub repository.
    The update pipeline clones the repo, plans the changes with Claude,
    generates new/modified file content, and commits + pushes to main.

    Requires GITHUB_TOKEN to be set with repo read/write access.
    """
    from app.api.main import get_build_queue
    from pipeline.update_pipeline import run_update_pipeline_sync

    update_id = str(uuid.uuid4())
    update = ForgeUpdate(
        update_id=update_id,
        repo_url=request.github_repo_url,
        change_description=request.change_description,
        title=request.title or f"Update: {request.github_repo_url.split('/')[-1]}",
        status="queued",
        callback_url=request.callback_url,
    )
    session.add(update)
    await session.commit()

    try:
        queue = get_build_queue()
        queue.enqueue(
            run_update_pipeline_sync,
            update_id,
            job_id=f"update-{update_id}",
            job_timeout=3600,
        )
        logger.info(
            f"Update queued: update_id={update_id} repo={request.github_repo_url}"
        )
    except Exception as exc:
        update.status = "failed"
        update.error_message = f"Failed to queue update: {exc}"
        await session.commit()
        logger.error(f"Failed to queue update {update_id}: {exc}")
        raise HTTPException(status_code=500, detail=f"Failed to queue update: {exc}")

    return UpdateResponse(
        update_id=update_id,
        status="queued",
        message="Update queued. Changes will be committed and pushed to main.",
    )


@router.get("/updates")
async def list_updates(
    session: AsyncSession = Depends(get_db),
) -> list[dict]:
    """List all update runs, most recent first."""
    from sqlalchemy import desc

    result = await session.execute(
        select(ForgeUpdate).order_by(desc(ForgeUpdate.created_at)).limit(100)
    )
    updates = result.scalars().all()
    return [
        {
            "update_id": u.update_id,
            "repo_url": u.repo_url,
            "title": u.title,
            "change_description": u.change_description[:200],
            "status": u.status,
            "files_created": u.files_created,
            "files_modified": u.files_modified,
            "files_deleted": u.files_deleted,
            "error_message": u.error_message,
            "created_at": u.created_at.isoformat() if u.created_at else None,
            "updated_at": u.updated_at.isoformat() if u.updated_at else None,
        }
        for u in updates
    ]


@router.get("/updates/{update_id}")
async def get_update(
    update_id: str,
    session: AsyncSession = Depends(get_db),
) -> dict:
    """Get full detail for a specific update run."""
    result = await session.execute(
        select(ForgeUpdate).where(ForgeUpdate.update_id == update_id)
    )
    update = result.scalar_one_or_none()
    if not update:
        raise HTTPException(status_code=404, detail="Update not found")

    return {
        "update_id": update.update_id,
        "repo_url": update.repo_url,
        "title": update.title,
        "change_description": update.change_description,
        "status": update.status,
        "files_created": update.files_created,
        "files_modified": update.files_modified,
        "files_deleted": update.files_deleted,
        "error_message": update.error_message,
        "changed_files_json": update.changed_files_json,
        "callback_url": update.callback_url,
        "created_at": update.created_at.isoformat() if update.created_at else None,
        "updated_at": update.updated_at.isoformat() if update.updated_at else None,
    }


# ── Helpers ───────────────────────────────────────────────────────────────────


def _slug(title: str) -> str:
    """Convert a title to a URL/repo-safe kebab-case slug."""
    slug = title.lower().strip()
    slug = re.sub(r"[^a-z0-9\s-]", "", slug)
    slug = re.sub(r"[\s_]+", "-", slug)
    slug = re.sub(r"-+", "-", slug)
    slug = slug.strip("-")
    return slug[:100] or "forge-build"


def _get_file_extension(filename: str) -> str:
    """Extract the lowercase file extension including the dot."""
    dot_index = filename.rfind(".")
    if dot_index == -1:
        return ""
    return filename[dot_index:].lower()


def _extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Extract plain text from uploaded file bytes.

    Accepts ANY file extension:
    - .docx: python-docx paragraph extraction
    - .pdf: pypdf page extraction
    - Known text extensions (.py, .ts, .go, etc.): UTF-8 decode
    - Unknown/no extension: attempt UTF-8, fall back to latin-1
    - If all parsing fails: raises ValueError (caller decides whether to skip or fail)

    Never raises due to an unrecognised extension — always attempts UTF-8 first.
    """
    ext = _get_file_extension(filename)

    # ── .docx extraction via python-docx ─────────────────────────────────
    if ext == ".docx":
        try:
            from docx import Document

            doc = Document(BytesIO(file_bytes))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            logger.info(
                f"Extracted {len(text)} chars from .docx file '{filename}' "
                f"({len(doc.paragraphs)} paragraphs)"
            )
            return text
        except ImportError:
            logger.warning("python-docx not installed — falling back to UTF-8 attempt")
        except Exception as exc:
            logger.error(f"Failed to extract text from .docx '{filename}': {exc}")
            raise ValueError(f"Failed to parse .docx file: {exc}") from exc

    # ── .pdf extraction via pypdf ────────────────────────────────────────
    if ext == ".pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(BytesIO(file_bytes))
            text = "\n".join(
                page.extract_text() or "" for page in reader.pages
            )
            logger.info(
                f"Extracted {len(text)} chars from .pdf file '{filename}' "
                f"({len(reader.pages)} pages)"
            )
            return text
        except ImportError:
            logger.warning("pypdf not installed — falling back to UTF-8 attempt")
        except Exception as exc:
            logger.error(f"Failed to extract text from .pdf '{filename}': {exc}")
            raise ValueError(f"Failed to parse .pdf file: {exc}") from exc

    # ── All other extensions (including unknown): attempt UTF-8 ──────────
    # This covers: .py .ts .go .rs .sh .sql .html .css .jsx .tsx .vue
    # .prisma .graphql .proto .cfg .ini .env .r .rb .php .swift .kt .dart
    # .svelte .makefile .dockerfile — and any unlisted or extension-less file.
    try:
        text = file_bytes.decode("utf-8")
        logger.info(
            f"Decoded {len(text)} chars from '{filename}' (UTF-8, ext='{ext or 'none'}')"
        )
        return text
    except UnicodeDecodeError:
        pass

    # UTF-8 failed — try latin-1 (never fails on any byte sequence)
    try:
        text = file_bytes.decode("latin-1")
        logger.warning(
            f"UTF-8 failed for '{filename}' — decoded {len(text)} chars via latin-1"
        )
        return text
    except Exception as exc:
        raise ValueError(
            f"Cannot decode file '{filename}' (ext='{ext or 'none'}') as text: {exc}"
        ) from exc


def _extract_docx_outline(full_text: str) -> str:
    """
    Extract a compact outline from DOCX text: headings + first 2 sentences per section.
    DOCX files are reference docs — the blueprint text already describes the content.
    This dramatically reduces token usage while preserving structure.
    """
    import re as _re

    lines = full_text.splitlines()
    result: list[str] = []
    i = 0

    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # Detect headings: ALL CAPS, short lines, or lines ending with colon
        is_heading = (
            (len(line) < 80 and line == line.upper() and len(line) > 3)
            or (len(line) < 60 and line.endswith(":"))
            or _re.match(r'^\d+[\.\)]\s+\S', line)
            or _re.match(r'^#{1,4}\s+\S', line)
        )

        if is_heading:
            result.append(line)
            # Include first 2 sentences of the following paragraph
            i += 1
            para_lines = []
            while i < len(lines) and lines[i].strip():
                para_lines.append(lines[i].strip())
                i += 1
            para = ' '.join(para_lines)
            # Split by sentence boundaries and take first 2
            sentences = _re.split(r'(?<=[.!?])\s+', para)
            if sentences:
                result.append(' '.join(sentences[:2]))
        else:
            i += 1

    outline = '\n'.join(result)
    if len(outline) < 100 and full_text:
        # Outline extraction failed — return a hard cap of the raw text
        return full_text[:3000]
    return outline