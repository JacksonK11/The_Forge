"""
app/api/services/file_extractor.py
Extracts text content from uploaded files based on file extension.

Supports:
  - python-docx for .docx files
  - pypdf for .pdf files
  - UTF-8 decode for text-based source and config files
  - Placeholder message for unsupported binary types

Used by the /forge/submit-with-files endpoint to process attached files
server-side before combining them with the blueprint text.
"""

import io
from pathlib import Path

from loguru import logger

# Extensions that are treated as plain UTF-8 text
TEXT_EXTENSIONS: set[str] = {
    ".py",
    ".txt",
    ".json",
    ".toml",
    ".md",
    ".yml",
    ".yaml",
    ".jsx",
    ".js",
    ".ts",
    ".tsx",
    ".css",
    ".html",
    ".xml",
    ".csv",
    ".cfg",
    ".ini",
    ".sh",
    ".bash",
    ".sql",
    ".r",
    ".rb",
    ".go",
    ".rs",
    ".java",
    ".c",
    ".cpp",
    ".h",
    ".env",
    ".gitignore",
    ".dockerignore",
    ".editorconfig",
    ".prettierrc",
    ".eslintrc",
    ".dockerfile",
    ".makefile",
    ".lock",
    ".log",
    ".rst",
    ".tex",
    ".bat",
    ".ps1",
    ".php",
    ".swift",
    ".kt",
    ".scala",
    ".lua",
    ".pl",
    ".pm",
    ".ex",
    ".exs",
    ".erl",
    ".hs",
    ".ml",
    ".clj",
    ".dart",
    ".v",
    ".vhdl",
    ".proto",
    ".graphql",
    ".tf",
    ".hcl",
}


def extract_text(filename: str, file_bytes: bytes) -> str:
    """
    Extract text content from a file based on its extension.

    Args:
        filename: Original filename (used to determine extension).
        file_bytes: Raw bytes of the uploaded file.

    Returns:
        Extracted text content, or a placeholder for unsupported binary files.
    """
    ext = Path(filename).suffix.lower()

    try:
        if ext == ".docx":
            return _extract_docx(filename, file_bytes)
        elif ext == ".pdf":
            return _extract_pdf(filename, file_bytes)
        elif ext in TEXT_EXTENSIONS or _looks_like_text(filename, ext):
            return _extract_text(filename, file_bytes)
        else:
            logger.info(f"Unsupported file type for extraction: {filename} (ext={ext})")
            return f"[Binary file: {filename} — content not extracted]"
    except Exception as exc:
        logger.error(f"Failed to extract text from {filename}: {exc}")
        return f"[Error extracting {filename}: {str(exc)}]"


def _extract_docx(filename: str, file_bytes: bytes) -> str:
    """Extract text from a .docx file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx is not installed — cannot extract .docx files")
        return f"[Cannot extract {filename}: python-docx not installed]"

    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts: list[str] = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))

        content = "\n\n".join(paragraphs)
        logger.info(f"Extracted {len(content)} chars from .docx: {filename}")
        return content
    except Exception as exc:
        logger.error(f"Error parsing .docx file {filename}: {exc}")
        raise


def _extract_pdf(filename: str, file_bytes: bytes) -> str:
    """Extract text from a .pdf file using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.error("pypdf is not installed — cannot extract .pdf files")
        return f"[Cannot extract {filename}: pypdf not installed]"

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        pages: list[str] = []

        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                pages.append(page_text.strip())

        content = "\n\n".join(pages)

        if not content.strip():
            logger.warning(
                f"PDF {filename} has {len(reader.pages)} pages but no extractable text "
                f"(possibly scanned/image-based)"
            )
            return f"[PDF file: {filename} — {len(reader.pages)} pages, no extractable text (scanned/image-based)]"

        logger.info(
            f"Extracted {len(content)} chars from {len(reader.pages)}-page PDF: {filename}"
        )
        return content
    except Exception as exc:
        logger.error(f"Error parsing .pdf file {filename}: {exc}")
        raise


def _extract_text(filename: str, file_bytes: bytes) -> str:
    """Extract text from a UTF-8 encoded text file."""
    try:
        content = file_bytes.decode("utf-8")
        logger.info(f"Extracted {len(content)} chars from text file: {filename}")
        return content
    except UnicodeDecodeError:
        # Try latin-1 as a fallback — it can decode any byte sequence
        try:
            content = file_bytes.decode("latin-1")
            logger.info(
                f"Extracted {len(content)} chars from text file (latin-1 fallback): {filename}"
            )
            return content
        except Exception:
            logger.warning(f"Could not decode {filename} as text")
            return f"[Binary file: {filename} — content not extracted]"


def _looks_like_text(filename: str, ext: str) -> bool:
    """
    Heuristic check for files that might be text but have unusual extensions.
    Checks if the filename has no extension (like Makefile, Dockerfile, Procfile)
    or matches known extensionless text files.
    """
    name_lower = Path(filename).name.lower()

    # Known extensionless text files
    known_text_files = {
        "makefile",
        "dockerfile",
        "procfile",
        "gemfile",
        "rakefile",
        "vagrantfile",
        "jenkinsfile",
        "cmakelists.txt",
        "license",
        "readme",
        "changelog",
        "authors",
        "contributing",
        "todo",
        "requirements",
    }

    if name_lower in known_text_files:
        return True

    # No extension — might be a text file
    if not ext:
        return True

    return False


def combine_blueprint_with_files(
    blueprint_text: str,
    extracted_files: list[tuple[str, str]],
) -> str:
    """
    Combine the user's blueprint text with extracted file contents.

    Args:
        blueprint_text: The user's typed instructions/prompt.
        extracted_files: List of (filename, extracted_content) tuples.

    Returns:
        Combined text with blueprint + all file contents separated by headers.
    """
    if not extracted_files:
        return blueprint_text

    parts: list[str] = [blueprint_text.strip()]

    for filename, content in extracted_files:
        parts.append(f"\n\n=== ATTACHED FILE: {filename} ===\n{content}")

    combined = "\n".join(parts)
    logger.info(
        f"Combined blueprint ({len(blueprint_text)} chars) with "
        f"{len(extracted_files)} attached file(s) → {len(combined)} chars total"
    )
    return combined